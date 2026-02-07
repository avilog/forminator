"""
Generate a test Word document with all supported field types.

Run:  python generate_test_doc.py
Output: demo_form.docx  (in the same folder)

After generating, open in Word, save as .docm, paste VBA, and run the macro.
The .txt files in this folder will be read automatically as context.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_content_control(paragraph, title, tag, cc_type="text", placeholder="Click here to enter text"):
    """Add a structured document tag (content control) inline."""
    run = paragraph.add_run()
    # Build the SDT XML
    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")

    # Tag
    tag_elem = OxmlElement("w:tag")
    tag_elem.set(qn("w:val"), tag)
    sdtPr.append(tag_elem)

    # Alias (title)
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), title)
    sdtPr.append(alias)

    # Placeholder
    ph = OxmlElement("w:placeholder")
    docPart = OxmlElement("w:docPart")
    docPart.set(qn("w:val"), placeholder)
    ph.append(docPart)
    sdtPr.append(ph)

    # Show-when-empty
    show = OxmlElement("w:showingPlcHdr")
    sdtPr.append(show)

    sdt.append(sdtPr)

    # SDT content
    sdtContent = OxmlElement("w:sdtContent")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "PlaceholderText")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = placeholder
    r.append(t)
    sdtContent.append(r)
    sdt.append(sdtContent)

    # Insert into paragraph
    paragraph._element.append(sdt)
    return sdt


def add_dropdown_control(paragraph, title, tag, options):
    """Add a dropdown content control."""
    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")

    tag_elem = OxmlElement("w:tag")
    tag_elem.set(qn("w:val"), tag)
    sdtPr.append(tag_elem)

    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), title)
    sdtPr.append(alias)

    # Dropdown list
    ddList = OxmlElement("w:dropDownList")
    for opt in options:
        li = OxmlElement("w:listItem")
        li.set(qn("w:displayText"), opt)
        li.set(qn("w:value"), opt)
        ddList.append(li)
    sdtPr.append(ddList)

    show = OxmlElement("w:showingPlcHdr")
    sdtPr.append(show)

    sdt.append(sdtPr)

    sdtContent = OxmlElement("w:sdtContent")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "PlaceholderText")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "Choose an item."
    r.append(t)
    sdtContent.append(r)
    sdt.append(sdtContent)

    paragraph._element.append(sdt)
    return sdt


def main():
    doc = Document()

    # ── Styles ──
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    # ── Title ──
    title = doc.add_heading("Employee Equipment Request Form", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Please fill out this form to request new or replacement equipment. "
        "All fields are required unless marked optional."
    )
    doc.add_paragraph("")

    # ═══════════════════════════════════════════════
    # Section 1: Content Controls
    # ═══════════════════════════════════════════════
    doc.add_heading("Section 1: Employee Information", level=2)

    p = doc.add_paragraph("Full Name: ")
    add_content_control(p, "Full Name", "FULL_NAME", placeholder="Enter full name")

    p = doc.add_paragraph("Employee ID: ")
    add_content_control(p, "Employee ID", "EMPLOYEE_ID", placeholder="Enter employee ID")

    p = doc.add_paragraph("Department: ")
    add_dropdown_control(p, "Department", "DEPARTMENT", [
        "Engineering", "Marketing", "Sales", "Human Resources",
        "Finance", "Operations", "Legal"
    ])

    p = doc.add_paragraph("Position: ")
    add_content_control(p, "Position", "POSITION", placeholder="Enter job title")

    p = doc.add_paragraph("Office Location: ")
    add_content_control(p, "Office Location", "OFFICE_LOCATION", placeholder="Building, floor, desk")

    p = doc.add_paragraph("Manager Name: ")
    add_content_control(p, "Manager Name", "MANAGER_NAME", placeholder="Enter manager name")

    doc.add_paragraph("")

    # ═══════════════════════════════════════════════
    # Section 2: Placeholders
    # ═══════════════════════════════════════════════
    doc.add_heading("Section 2: Request Details", level=2)

    doc.add_paragraph(
        "I, <<FULL_NAME>>, employee ID <<EMPLOYEE_ID>>, am requesting "
        "the following equipment for use at <<OFFICE_LOCATION>>."
    )

    doc.add_paragraph(
        "This request is submitted on <<DATE>> and should be approved by "
        "<<MANAGER_NAME>> from the <<DEPARTMENT>> department."
    )

    doc.add_paragraph("")

    # ═══════════════════════════════════════════════
    # Section 3: Underscore blanks
    # ═══════════════════════════════════════════════
    doc.add_heading("Section 3: Equipment Specifications", level=2)

    doc.add_paragraph("Equipment type requested: ________________________")
    doc.add_paragraph("Reason for request: ________________________________________")
    doc.add_paragraph("Preferred brand/model (optional): ________________________")
    doc.add_paragraph("Budget estimate: $________________________")

    doc.add_paragraph("")

    # ═══════════════════════════════════════════════
    # Section 4: Checkboxes
    # ═══════════════════════════════════════════════
    doc.add_heading("Section 4: Request Type", level=2)

    doc.add_paragraph("This is a:  [ ] New equipment request  [ ] Replacement for damaged/lost item  [ ] Upgrade request")

    doc.add_paragraph("Priority:  [ ] Low (within 30 days)  [ ] Medium (within 2 weeks)  [ ] High (within 3 days)  [ ] Urgent (next business day)")

    doc.add_paragraph("")

    doc.add_heading("Section 5: Acknowledgments", level=2)

    doc.add_paragraph("[ ] I confirm the information provided is accurate")
    doc.add_paragraph("[ ] I agree to the company equipment usage policy")
    doc.add_paragraph("[ ] I understand that approval is subject to budget availability")

    doc.add_paragraph("")

    # ═══════════════════════════════════════════════
    # Section 6: Some prose (for comment-based revision)
    # ═══════════════════════════════════════════════
    doc.add_heading("Section 6: Justification", level=2)

    doc.add_paragraph(
        "The employee need the new equipment becuz their current laptop is very slow and "
        "it crashes alot when running development tools. This is make it hard to do work "
        "on the Atlas Migration project which is high priority. Getting new equipment will "
        "help them be more productive and meet the project deadlines."
    )

    doc.add_paragraph(
        "The team is currently in Phase 2 of the Atlas Migration project and the delays "
        "caused by hardware issues are affecting the whole team's velocity. The estimated "
        "ROI of this equipment upgrade is significant given the project's budget and timeline."
    )

    doc.add_paragraph("")

    # ── Signature area ──
    doc.add_heading("Signatures", level=2)
    doc.add_paragraph("Employee signature: ________________________    Date: ____________")
    doc.add_paragraph("Manager signature:  ________________________    Date: ____________")

    # ── Save ──
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "demo_form.docx")
    doc.save(out_path)
    print(f"Generated: {out_path}")
    print()
    print("Next steps:")
    print("  1. Open demo_form.docx in Word")
    print("  2. Save As -> demo_form.docm (macro-enabled)")
    print("  3. Alt+F11 -> Insert -> Module -> paste word.vba")
    print("  4. Add a comment on the justification paragraph:")
    print('     e.g. "fix grammar and make more professional"')
    print("  5. Run the macro: Alt+F8 -> FillForm_Tracked_ByPython")


if __name__ == "__main__":
    main()

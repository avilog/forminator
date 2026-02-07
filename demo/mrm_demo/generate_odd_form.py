"""
Generate a simple MRM Object Development Document (ODD) form for demo.

Run:  python generate_odd_form.py
Output: odd_form.docx  (in the same folder)

After generating, open in Word, save as .docm, paste VBA, and run the macro.
The model_info.txt file in this folder provides context for the LLM.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def add_content_control(paragraph, title, tag, placeholder="Click here to enter text"):
    """Add a structured document tag (content control) inline."""
    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")

    tag_elem = OxmlElement("w:tag")
    tag_elem.set(qn("w:val"), tag)
    sdtPr.append(tag_elem)

    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), title)
    sdtPr.append(alias)

    ph = OxmlElement("w:placeholder")
    docPart = OxmlElement("w:docPart")
    docPart.set(qn("w:val"), placeholder)
    ph.append(docPart)
    sdtPr.append(ph)

    show = OxmlElement("w:showingPlcHdr")
    sdtPr.append(show)
    sdt.append(sdtPr)

    sdtContent = OxmlElement("w:sdtContent")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "PlaceholderText")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = placeholder
    r.append(t)
    sdtContent.append(r)
    sdt.append(sdtContent)

    paragraph._element.append(sdt)
    return sdt


def add_dropdown_control(paragraph, title, tag, options):
    """Add a dropdown content control."""
    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")

    tag_elem = OxmlElement("w:tag")
    tag_elem.set(qn("w:val"), tag)
    sdtPr.append(tag_elem)

    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), title)
    sdtPr.append(alias)

    ddList = OxmlElement("w:dropDownList")
    for opt in options:
        li = OxmlElement("w:listItem")
        li.set(qn("w:displayText"), opt)
        li.set(qn("w:value"), opt)
        ddList.append(li)
    sdtPr.append(ddList)

    show = OxmlElement("w:showingPlcHdr")
    sdtPr.append(show)
    sdt.append(sdtPr)

    sdtContent = OxmlElement("w:sdtContent")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "PlaceholderText")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "Choose an item."
    r.append(t)
    sdtContent.append(r)
    sdt.append(sdtContent)

    paragraph._element.append(sdt)
    return sdt


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    # ── Title ──
    title = doc.add_heading("Object Development Document (ODD)", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Model Risk Management")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].italic = True

    doc.add_paragraph("")

    # ═══════════════════════════════════════════════
    # Section 1: Model Identification (Content Controls)
    # ═══════════════════════════════════════════════
    doc.add_heading("1. Model Identification", level=2)

    p = doc.add_paragraph("Model Name: ")
    add_content_control(p, "Model Name", "MODEL_NAME", placeholder="Enter model name")

    p = doc.add_paragraph("Model ID: ")
    add_content_control(p, "Model ID", "MODEL_ID", placeholder="Enter model ID")

    p = doc.add_paragraph("Model Tier: ")
    add_dropdown_control(p, "Model Tier", "MODEL_TIER", [
        "Tier 1 - Critical",
        "Tier 2 - Significant",
        "Tier 3 - Limited",
    ])

    p = doc.add_paragraph("Model Owner: ")
    add_content_control(p, "Model Owner", "MODEL_OWNER", placeholder="Enter model owner")

    p = doc.add_paragraph("Lead Developer: ")
    add_content_control(p, "Lead Developer", "LEAD_DEVELOPER", placeholder="Enter lead developer")

    p = doc.add_paragraph("Development Date: ")
    add_content_control(p, "Development Date", "DEV_DATE", placeholder="Enter date")

    doc.add_paragraph("")

    # ═══════════════════════════════════════════════
    # Section 2: Purpose & Scope (Placeholders)
    # ═══════════════════════════════════════════════
    doc.add_heading("2. Purpose & Scope", level=2)

    doc.add_paragraph(
        "The <<MODEL_NAME>> (ID: <<MODEL_ID>>) was developed by the "
        "<<LEAD_DEVELOPER>> team to support regulatory stress testing and "
        "capital planning requirements. The model is owned by <<MODEL_OWNER>> "
        "and was last updated on <<DEV_DATE>>."
    )

    doc.add_paragraph("")

    # ═══════════════════════════════════════════════
    # Section 3: Data & Methodology (Underscore blanks)
    # ═══════════════════════════════════════════════
    doc.add_heading("3. Data & Methodology", level=2)

    doc.add_paragraph("Primary data source: ________________________________________")
    doc.add_paragraph("Observation period: ________________________")
    doc.add_paragraph("Key methodology: ________________________________________")
    doc.add_paragraph("Number of risk factors: ________________________")

    doc.add_paragraph("")

    # ═══════════════════════════════════════════════
    # Section 4: Model Classification (Checkboxes)
    # ═══════════════════════════════════════════════
    doc.add_heading("4. Model Classification", level=2)

    doc.add_paragraph(
        "Model Use:  [ ] Regulatory reporting  [ ] Internal risk management  [ ] Pricing  [ ] Client-facing"
    )

    doc.add_paragraph(
        "Risk Rating:  [ ] High  [ ] Medium  [ ] Low"
    )

    doc.add_paragraph("")

    # ═══════════════════════════════════════════════
    # Section 5: Executive Summary (prose with errors - for revision demo)
    # ═══════════════════════════════════════════════
    doc.add_heading("5. Executive Summary", level=2)

    doc.add_paragraph(
        "The model use a monte carlo simulation approach to estimate potential "
        "lossess under various stress scenarios. It take historical market data "
        "and apply stress factors to generate forward-looking loss distributions. "
        "The model have been validated against the 2008 financial crisis data and "
        "recent market stress events, demostrating strong predictive performance."
    )

    doc.add_paragraph("")

    # ── Sign-off ──
    doc.add_heading("Sign-off", level=2)
    doc.add_paragraph("Model Developer: ________________________    Date: ____________")
    doc.add_paragraph("Model Owner:     ________________________    Date: ____________")
    doc.add_paragraph("MRM Reviewer:    ________________________    Date: ____________")

    # ── Save ──
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "odd_form.docx")
    doc.save(out_path)
    print(f"Generated: {out_path}")
    print()
    print("Next steps:")
    print("  1. Open odd_form.docx in Word")
    print("  2. Save As -> odd_form.docm (macro-enabled)")
    print("  3. Alt+F11 -> Insert -> Module -> paste word.vba")
    print("  4. Run the macro: Alt+F8 -> FillForm_Tracked_ByPython")
    print("  5. After filling, add a comment on the Executive Summary:")
    print('     "fix grammar and make more professional"')
    print("  6. Run the macro again to apply the fix")


if __name__ == "__main__":
    main()
